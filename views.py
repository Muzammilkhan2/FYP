from django.shortcuts import render
from django.conf import settings
from django.utils.decorators import method_decorator
from datetime import datetime
from urllib.parse import urlparse, parse_qs
from .models import * 
from .serializers import *
import tiktoken
import nltk
from rest_framework import viewsets
from nltk.corpus import stopwords
from nltk.tokenize import sent_tokenize, word_tokenize
from rest_framework.decorators import api_view, permission_classes, APIView, action
from rest_framework.response import Response
from rest_framework.permissions import AllowAny
from io import BytesIO
from youtube_transcript_api._errors import TranscriptsDisabled, NoTranscriptFound
from django.template.loader import get_template
import textwrap
from django.core.exceptions import ObjectDoesNotExist
from django.http import HttpResponse
from django.db.models import Avg, Count, Max, Min
from reportlab.lib.pagesizes import letter
from reportlab.lib.pagesizes import A4
from django.db.models import Avg, Count, Q
from reportlab.lib.units import inch
from xhtml2pdf import pisa
from urllib.parse import urlparse
from rest_framework import status, generics, permissions
from rest_framework_simplejwt.tokens import RefreshToken
from reportlab.pdfgen import canvas
from rest_framework.generics import ListAPIView
from rest_framework_simplejwt.views import TokenObtainPairView
from rest_framework_simplejwt.serializers import TokenObtainPairSerializer
from django.contrib.auth import authenticate
from rest_framework.permissions import IsAuthenticated
from django.http import JsonResponse
from youtube_transcript_api import YouTubeTranscriptApi
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth import get_user_model
from rest_framework.permissions import BasePermission
from django.shortcuts import get_object_or_404
from rest_framework.parsers import MultiPartParser, FormParser
from bs4 import BeautifulSoup
import yt_dlp
from django.core.files.storage import default_storage
import fitz  # PyMuPDF for PDF
from utils.notification import send_notification
from rest_framework.generics import RetrieveAPIView
import docx 
import requests
import json
import re
import random
from openai import OpenAI
import logging
import asyncio
from utils.AI_Model import evaluate_with_AI
from django.db import transaction
import time
import aiohttp
import traceback
import io

logger = logging.getLogger(__name__)

STOP_WORDS = set(stopwords.words('english'))

User = get_user_model()

class RegisterView(generics.CreateAPIView):
    queryset = User.objects.all()
    serializer_class = RegisterSerializer
    permission_classes = [AllowAny]
    

class LoginView(APIView):
    permission_classes = [AllowAny]

    def post(self, request, *args, **kwargs):
        email = request.data.get("email")
        password = request.data.get("password")

        user = User.objects.filter(email=email).first()

        if user and user.check_password(password):
            if user.role == "teacher" and not user.is_verified:
                return Response({"message": "Admin approval required"}, status=status.HTTP_403_FORBIDDEN)

            # Create the JWT with the role field
            refresh = RefreshToken.for_user(user)
            refresh.payload['role'] = user.role  # Add role to the token payload

            return Response({
                "refresh": str(refresh),
                "access": str(refresh.access_token),
                "user": {
                    "id": user.id,
                    "email": user.email,
                    "role": user.role,  # Send the role in the response as well
                    "is_verified": user.is_verified
                }
            })

        return Response({"error": "Invalid credentials"}, status=status.HTTP_401_UNAUTHORIZED)



class IsCustomAdmin(BasePermission):
    def has_permission(self, request, view):
        return request.user.is_authenticated and request.user.role == "admin"

class AdminOnlyView(APIView):
    permission_classes = [IsAuthenticated, IsCustomAdmin]

    def get(self, request):
        return Response({"message": "Hello, Admin!"})


class ApproveTeacherView(APIView):
    permission_classes = [IsCustomAdmin]

    def patch(self, request, teacher_id):
        teacher = get_object_or_404(User, id=teacher_id, role="teacher")
        teacher.is_verified = True
        teacher.save()
        return Response({"message": "Teacher approved successfully"}, status=status.HTTP_200_OK)

@api_view(["GET"])
@permission_classes([IsAuthenticated])
def get_unverified_teachers(request):
    print("Headers Received:", request.headers)  # Debugging
    if request.user.role != "admin":
        return Response({"error": "Unauthorized"}, status=status.HTTP_403_FORBIDDEN)

    teachers = User.objects.filter(role="teacher", is_verified=False)
    data = [{"id": teacher.id, "email": teacher.email} for teacher in teachers]
    return Response(data, status=status.HTTP_200_OK)


class CreateSessionView(generics.CreateAPIView):
    serializer_class = SessionSerializer
    permission_classes = [permissions.IsAuthenticated]

    def perform_create(self, serializer):
        if self.request.user.role != "teacher":
            raise serializers.ValidationError("Only teachers can create sessions.")
        serializer.save(teacher=self.request.user)

# Endpoint to list all sessions (for enrollment)
class ListSessionsView(generics.ListAPIView):
    serializer_class = SessionSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        return Session.objects.all()

# Endpoint for students to enroll in a session
class EnrollSessionView(generics.UpdateAPIView):
    serializer_class = SessionSerializer
    permission_classes = [permissions.IsAuthenticated]

    def update(self, request, *args, **kwargs):
        session = self.get_object()
        if request.user.role != "student":
            return Response({"error": "Only students can enroll in sessions."}, status=status.HTTP_403_FORBIDDEN)
        session.enrolled_students.add(request.user)
        return Response({"message": "Enrolled successfully."}, status=status.HTTP_200_OK)
    
    def get_queryset(self):
        return Session.objects.all()
    
class RequestEnrollmentView(generics.UpdateAPIView):
    serializer_class = SessionSerializer
    permission_classes = [permissions.IsAuthenticated]

    def update(self, request, *args, **kwargs):
        session = self.get_object()
        if request.user.role != "student":
            return Response({"error": "Only students can request enrollment."}, status=status.HTTP_403_FORBIDDEN)
        if request.user in session.pending_students.all():
            return Response({"error": "Already requested enrollment."}, status=status.HTTP_400_BAD_REQUEST)
        session.pending_students.add(request.user)
        return Response({"message": "Enrollment request sent."}, status=status.HTTP_200_OK)

    def get_queryset(self):
        return Session.objects.all()
    

class ManageEnrollmentsView(generics.UpdateAPIView):
    serializer_class = SessionSerializer
    permission_classes = [permissions.IsAuthenticated]
    queryset = Session.objects.all()
    lookup_field = 'pk'  # <-- Add this line!

    def update(self, request, *args, **kwargs):
        session = self.get_object()

        if request.user.role != "teacher":
            return Response({"error": "Only verified teachers can manage enrollments."}, status=status.HTTP_403_FORBIDDEN)
        
        student_id = request.data.get("student_id")
        action = request.data.get("action")

        student = session.pending_students.filter(id=student_id).first()
        if not student:
            return Response({"error": "Student not found in pending list."}, status=status.HTTP_404_NOT_FOUND)

        if action == "approve":
            session.pending_students.remove(student)
            session.enrolled_students.add(student)
            return Response({"message": "Student approved."}, status=status.HTTP_200_OK)
        elif action == "reject":
            session.pending_students.remove(student)
            return Response({"message": "Student rejected."}, status=status.HTTP_200_OK)
        else:
            return Response({"error": "Invalid action."}, status=status.HTTP_400_BAD_REQUEST)

class TeacherSessionsView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        print("User:", request.user)
        print("User type:", type(request.user))
        print("Role:", getattr(request.user, 'role', 'No role'))

        if getattr(request.user, 'role', None) != "teacher":
            return Response({"error": "Only teachers can access their sessions."}, status=403)

        sessions = Session.objects.filter(teacher=request.user)
        serializer = SessionSerializer(sessions, many=True)
        return Response(serializer.data)

# views.py
class CreateTestView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request, session_id):
        if request.user.role != "teacher":
            return Response({"error": "Only teachers can create tests."}, status=403)

        try:
            session = Session.objects.get(id=session_id, teacher=request.user)
        except Session.DoesNotExist:
            return Response({"error": "Session not found or not owned by you."}, status=404)

        data = request.data.copy()
        data['teacher'] = request.user.id
        data['session'] = session.id

        serializer = TestSerializer(data=data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=201)
        return Response(serializer.errors, status=400)


# Ensure NLTK resources are downloaded once
def ensure_nltk_resources():
    resources = {
        'punkt': ('tokenizers/punkt', nltk.download),
        'stopwords': ('corpora/stopwords', nltk.download),
        'wordnet': ('corpora/wordnet', nltk.download),
        'omw-1.4': ('corpora/omw-1.4', nltk.download)
    }
    
    for resource, (path, downloader) in resources.items():
        try:
            nltk.data.find(path)
        except LookupError:
            downloader(resource.split('/')[-1], quiet=True)

class GenerateQuestionsView(APIView):
    permission_classes = [IsAuthenticated]
    parser_classes = [MultiPartParser, FormParser]
    
    # Configuration
    MAX_PDF_SIZE = 10 * 1024 * 1024  # 10MB
    MAX_CONTENT_TOKENS = 80000
    SUMMARY_RATIO = 0.3
    MODEL_MAX_TOKENS = 8000

    def __init__(self):
        self.tokenizer = tiktoken.get_encoding("cl100k_base")
        self.token_counts = {}
        ensure_nltk_resources()

    def post(self, request):
        """Main endpoint for question generation"""
        try:
            # Extract and process content
            raw_content = self._extract_content(request)
            processed_content = self._process_content(raw_content)
            difficulty = request.data.get('difficulty', 'medium').capitalize()
            
            # Create AI prompt
            prompt, prompt_tokens = self._create_prompt(
                processed_content,
                int(request.data.get('mcq_count', 0)),
                int(request.data.get('qna_count', 0)),
                difficulty
            )
            
            # Generate and format questions
            raw_questions = self._generate_questions(prompt)
            formatted_questions = self._format_questions(
                raw_questions, 
                difficulty=difficulty,
                topic=processed_content[:100]  # Truncate to match model's max_length
            )
            
            return Response({
                "questions": formatted_questions,
                "token_usage": {
                    "prompt_tokens": prompt_tokens,
                    "content_tokens": self.token_counts.get('content', 0),
                    "total_tokens": prompt_tokens + self.token_counts.get('content', 0)
                },
                "warning": "Questions are not persisted to database"
            })

        except ValueError as e:
            logger.error(f"Validation Error: {str(e)}")
            return Response({"error": str(e)}, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            logger.error(f"Server Error: {traceback.format_exc()}")
            return Response(
                {"error": "Internal server error. Please try again later."},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
            

    def _extract_content(self, request):
        """Extract content from file/url/text input"""
        if file := request.FILES.get('file'):
            if file.size > self.MAX_PDF_SIZE:
                raise ValueError(f"File size exceeds {self.MAX_PDF_SIZE//1024//1024}MB limit")
            return self._process_file(file)
        
        text = request.data.get('prompt_text', '')
        if text.startswith(('http://', 'https://')):
            return self._process_webpage(text)
            
        return text

    def _process_file(self, file):
        """Process uploaded files in memory"""
        try:
            if file.name.endswith('.pdf'):
                with fitz.open(stream=file.read(), filetype="pdf") as doc:
                    return " ".join(page.get_text() for page in doc)
            elif file.name.endswith('.docx'):
                return " ".join(p.text for p in docx.Document(io.BytesIO(file.read())).paragraphs)
            elif file.name.endswith('.txt'):
                return file.read().decode('utf-8')
            raise ValueError("Unsupported file format")
        except Exception as e:
            raise ValueError(f"File processing error: {str(e)}")

    def _process_webpage(self, url):
        """Extract and clean webpage content"""
        try:
            response = requests.get(url, timeout=10)
            response.raise_for_status()
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Remove unwanted elements
            for tag in soup(['script', 'style', 'nav', 'footer', 'aside', 'header']):
                tag.decompose()
                
            return ' '.join(soup.stripped_strings)[:100000] 
        except Exception as e:
            raise ValueError(f"Webpage processing failed: {str(e)}")

    def _process_content(self, text):
        """Ensure content token count stays under 80k using NLTK summarization."""
        tokens = self.tokenizer.encode(text)
        self.token_counts['content'] = len(tokens)

        if len(tokens) <= self.MAX_CONTENT_TOKENS:
            return text

        # Reduce the content using iterative summarization
        reduced_text = text
        max_iterations = 5  # prevent infinite loops
        ratio = self.SUMMARY_RATIO

        for _ in range(max_iterations):
            reduced_text = self._summarize_text(reduced_text, ratio=ratio)
            tokens = self.tokenizer.encode(reduced_text)
            self.token_counts['content'] = len(tokens)
            if len(tokens) <= self.MAX_CONTENT_TOKENS:
                break
            ratio *= 0.7  # progressively reduce more

        if len(tokens) > self.MAX_CONTENT_TOKENS:
            raise ValueError("Input content is too long even after summarization")

        return reduced_text


    def _summarize_text(self, text, ratio=None):
        """Text summarization using sentence scoring with adjustable ratio"""
        ratio = ratio if ratio is not None else self.SUMMARY_RATIO
        sentences = sent_tokenize(text)
        if not sentences:
            return text

        words = [word.lower() for word in word_tokenize(text) 
                if word.isalnum() and word not in stopwords.words('english')]
        
        freq_dist = nltk.FreqDist(words)
        sentence_scores = {}

        for i, sentence in enumerate(sentences):
            for word in word_tokenize(sentence.lower()):
                if word in freq_dist:
                    sentence_scores[i] = sentence_scores.get(i, 0) + freq_dist[word]

        keep_count = max(1, int(len(sentences) * ratio))
        top_sentences = sorted(sentence_scores.items(), key=lambda x: x[1], reverse=True)[:keep_count]
        top_sentences = sorted([s[0] for s in top_sentences])

        return ' '.join([sentences[i] for i in top_sentences])


    def _create_prompt(self, content, mcq_count, qna_count, difficulty):
        """Construct and tokenize the AI prompt with strict format requirements"""
        prompt_template = f"""
    [STRICTLY FOLLOW THESE FORMAT RULES]
    - Output MUST be pure JSON only (no markdown/code blocks)
    - Top-level keys MUST be "mcqs" and "qnas"
    - MCQ format:
        {{
          "content": "Question text",
          "option_a": "Choice A",
          "option_b": "Choice B",
          "option_c": "Choice C",
          "option_d": "Choice D",
          "correct_option": "b"  // Must be a/b/c/d
        }}
    - QnA format:
        {{ "content": "Question text" }}

    [EXAMPLE]
    {{
      "mcqs": [{{...}}],
      "qnas": [{{...}}]
    }}

    Generate for:
    - Topic: {content[:1000]}... [truncated]
    - Difficulty: {difficulty.upper()}
    - MCQs: {mcq_count}, QnAs: {qna_count}
    """

        # Tokenization logic remains the same
        tokens = self.tokenizer.encode(prompt_template)
        if len(tokens) > self.MODEL_MAX_TOKENS:
            truncated = self.tokenizer.decode(tokens[:self.MODEL_MAX_TOKENS])
            return truncated, self.MODEL_MAX_TOKENS
        return prompt_template, len(tokens)


    def _extract_json(self, content):
        """Robust JSON extraction with validation"""
        try:
            # Handle code block formatting
            if "```json" in content:
                content = content.split("```json")[1].split("```")[0]
            elif "```" in content:
                content = content.split("```")[1].split("```")[0]

            # Attempt direct parse
            data = json.loads(content.strip())
            
            if 'quiz' in data and 'questions' in data['quiz']:
                return self._remap_quiz_structure(data['quiz']['questions'])
            
            # Validate structure
            if not isinstance(data, dict):
                raise ValueError("Top-level structure is not a dictionary")
                
            if 'mcqs' not in data or 'qnas' not in data:
                raise ValueError("Missing required 'mcqs' or 'qnas' keys")
                
            return data

        except json.JSONDecodeError as e:
            logger.warning(f"JSON Decode Error: {e}")
            # Fallback: Try to find JSON substring
            try:
                json_str = re.search(r'\{.*\}', content, re.DOTALL)
                if json_str:
                    return json.loads(json_str.group())
                raise
            except Exception as e:
                logger.error(f"Fallback JSON parse failed: {e}")
                raise ValueError("Could not extract valid JSON from response")
            
    def _remap_quiz_structure(self, questions):
        """Convert 'questions' array to {mcqs: [], qnas: []} format"""
        mcqs = []
        qnas = []
        
        for q in questions:
            if 'options' in q:  # MCQ
                mcqs.append({
                    "content": q.get("question", ""),
                    "option_a": q['options'][0] if len(q['options']) > 0 else "",
                    "option_b": q['options'][1] if len(q['options']) > 1 else "",
                    "option_c": q['options'][2] if len(q['options']) > 2 else "",
                    "option_d": q['options'][3] if len(q['options']) > 3 else "",
                    "correct_option": self._derive_correct_option(q)
                })
            else:  # QnA
                qnas.append({"content": q.get("question", "")})
        
        return {"mcqs": mcqs, "qnas": qnas}

    def _derive_correct_option(self, question):
        """Convert answer text to A/B/C/D index"""
        answer = str(question.get("answer", "")).strip()
        options = question.get("options", [])
        
        # Find matching option index
        for idx, opt in enumerate(options):
            if str(opt).strip() == answer:
                return chr(65 + idx)  # 65 = 'A'
        
        return "A"  # Default if not found

    def _generate_questions(self, prompt):
        """Handle AI response with multiple fallback strategies"""
        try:
            ai_response = evaluate_with_AI(prompt)
            if not ai_response:
                raise ValueError("Empty response from AI model")

            logger.debug(f"Raw AI Response: {ai_response[:500]}...")  # Log first 500 chars

            # First try: Structured JSON parse
            try:
                parsed_data = self._extract_json(ai_response)
                return parsed_data
            except ValueError as e:
                logger.warning(f"Primary JSON parse failed: {e}")

            # Fallback 1: Try direct JSON parse
            try:
                return json.loads(ai_response)
            except json.JSONDecodeError:
                pass

            # Fallback 2: Manual content parsing
            try:
                return self._extract_questions_from_content(ai_response)
            except Exception as e:
                logger.warning(f"Content parsing failed: {e}")

            # Final fallback: Return raw response
            logger.error("All parsing strategies failed")
            return {"mcqs": [], "qnas": []}

        except Exception as e:
            logger.error(f"Question generation failed: {traceback.format_exc()}")
            raise ValueError("Failed to generate questions. Please try again.")

    def _format_questions(self, raw_data, difficulty, topic):
        """Convert raw AI response to model-compatible format"""
        formatted = []
        
        # Process MCQs
        for mcq in raw_data.get('mcqs', []):
            if not all(k in mcq for k in ["content", "option_a", "option_b", 
                    "option_c", "option_d", "correct_option"]):
                logger.warning(f"Skipping invalid MCQ: {mcq}")
                continue
            try:
                formatted.append({
                    "question_type": "MCQ",
                    "content": mcq['content'],
                    "option_a": mcq.get('option_a', ''),
                    "option_b": mcq.get('option_b', ''),
                    "option_c": mcq.get('option_c', ''),
                    "option_d": mcq.get('option_d', ''),
                    "correct_option": mcq.get('correct_option', '').upper(),
                    "difficulty": difficulty,
                    "topic": topic
                })
            except KeyError as e:
                logger.warning(f"Skipping invalid MCQ: Missing {str(e)}")

        # Process QNAs
        for qna in raw_data.get('qnas', []):
            try:
                formatted.append({
                    "question_type": "QNA",
                    "content": qna['content'],
                    "option_a": None,
                    "option_b": None,
                    "option_c": None,
                    "option_d": None,
                    "correct_option": None,
                    "difficulty": difficulty,
                    "topic": topic
                })
            except KeyError as e:
                logger.warning(f"Skipping invalid QNA: Missing {str(e)}")

        if not formatted:
            raise ValueError("No valid questions could be extracted")
            
        return formatted


    def _extract_questions_from_content(self, content):
        """Fallback parser for non-JSON responses"""
        questions = {"mcqs": [], "qnas": []}
        
        # Improved regex patterns
        mcq_pattern = r"\d+[\.\)]\s*(.+?)\n(a\)\s*.+?\n)b\)\s*(.+?)\nc\)\s*(.+?)\nd\)\s*(.+?)\n.*Answer:\s*([A-D])"
        qna_pattern = r"\d+[\.\)]\s*(.+?)(?:\nAnswer:\s*(.+?))(?=\n\d+|\Z)"
        
        # MCQs
        for match in re.finditer(mcq_pattern, content, re.DOTALL | re.IGNORECASE):
            try:
                question, a, b, c, d, answer = match.groups()
                questions["mcqs"].append({
                    "content": question.strip(),
                    "option_a": a.strip(),
                    "option_b": b.strip(),
                    "option_c": c.strip(),
                    "option_d": d.strip(),
                    "correct_option": answer.upper()
                })
            except Exception as e:
                logger.warning(f"Failed to parse MCQ: {str(e)}")

        # QNAs
        for match in re.finditer(qna_pattern, content, re.DOTALL):
            try:
                question, answer = match.groups()
                questions["qnas"].append({
                    "content": question.strip(),
                    "answer": answer.strip() if answer else ""
                })
            except Exception as e:
                logger.warning(f"Failed to parse QNA: {str(e)}")

        return questions


class SaveQuizView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        questions = request.data.get('questions', [])
        test_id = request.data.get('test_id')

        try:
            test = Test.objects.get(id=test_id)
        except ObjectDoesNotExist:
            return Response({'error': 'Test not found'}, status=404)

        saved_questions = []
        errors = []

        for item in questions:
            # Skip 'id' field handling if it's invalid
            question_id = item.get('id', None)
            question_type = item.get('question_type', 'MCQ')

            base_data = {
                'question_type': question_type,
                'content': item.get('content', '').strip(),
                'difficulty': item.get('difficulty', 'Medium'),
                'topic': item.get('topic'),
                'test': test,
                'teacher': request.user
            }

            # MCQ-specific fields
            if question_type == 'MCQ':
                base_data.update({
                    'option_a': item.get('option_a'),
                    'option_b': item.get('option_b'),
                    'option_c': item.get('option_c'),
                    'option_d': item.get('option_d'),
                    'correct_option': item.get('correct_option'),
                })

            try:
                if question_id and question_id.startswith("temp-"):
                    # Handle temp ids - create new questions instead of updating
                    question = Question.objects.create(**base_data)
                else:
                    # Handle actual existing question IDs
                    if question_id:
                        question = Question.objects.get(id=question_id)
                        for key, value in base_data.items():
                            setattr(question, key, value)
                        question.save()
                    else:
                        # Create a new question if no question_id is provided
                        question = Question.objects.create(**base_data)

                saved_questions.append({
                    'id': question.id,
                    'content': question.content,
                    'question_type': question.question_type,
                })

            except Exception as e:
                print("❌ Error saving question:", item)
                print("❌ Exception:", str(e))
                errors.append({
                    'question': item,
                    'error': str(e),
                })

        return Response({
            'saved': saved_questions,
            'errors': errors
        })

            
# views.py
class SetTimeLimitView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, test_id):
        test = get_object_or_404(Test, id=test_id, teacher=request.user)
        time_limit = request.data.get("time_limit_minutes")
        
        if not time_limit or time_limit <= 0:
            return Response({"error": "Invalid time limit"}, status=400)
        
        test.time_limit_minutes = time_limit
        test.save()
        return Response({"message": f"Time limit set to {time_limit} minutes"})
    
class EnrolledSessionsView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        # Grab all Session objects where the current user is enrolled
        sessions = request.user.enrolled_sessions.all()
        serializer = SessionSerializer(sessions, many=True)
        return Response(serializer.data, status=200)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def enrolled_sessions_with_tests(request):
    student = request.user
    sessions = Session.objects.filter(enrolled_students=student).distinct()
    serializer = SessionWithTestsSerializer(sessions, many=True)
    return Response(serializer.data)

class StudentSessionTestsView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, session_id):
        try:
            session = Session.objects.get(id=session_id)

            # Check if student is enrolled in the session
            if request.user not in session.enrolled_students.all():
                return Response({"detail": "You are not enrolled in this session."}, status=status.HTTP_403_FORBIDDEN)

            tests = Test.objects.filter(session=session)
            serializer = TestSerializer(tests, many=True)
            return Response(serializer.data)

        except Session.DoesNotExist:
            return Response({"detail": "Session not found."}, status=status.HTTP_404_NOT_FOUND)
        

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def get_test_for_attempt(request, test_id):
    try:
        # Get test or return 404
        test = get_object_or_404(Test, id=test_id)
        user = request.user

        # Validate user role
        if user.role != 'student':
            return Response(
                {"error": "Only students can attempt tests"},
                status=status.HTTP_403_FORBIDDEN
            )

        # Check enrollment
        if not test.session.enrolled_students.filter(id=user.id).exists():
            return Response(
                {"error": "You are not enrolled in this session"},
                status=status.HTTP_403_FORBIDDEN
            )

        # Create or get test attempt
        attempt, created = TestAttempt.objects.get_or_create(
            test=test,
            student=user,
            is_submitted=False,
            defaults={'start_time': timezone.now()}
        )

        # Prepare response data
        response_data = {
            "test_id": test.id,
            "title": test.title,
            "time_limit_minutes": test.time_limit_minutes,
            "attempt_id": attempt.id,
            "questions": []
        }

        # Add questions
        for question in test.questions.all():
            question_data = {
                "id": question.id,
                "content": question.content,
                "question_type": question.question_type,
                "marks": 1  # Default value
            }

            if question.question_type == 'MCQ':
                question_data['options'] = {
                    'A': question.option_a,
                    'B': question.option_b,
                    'C': question.option_c,
                    'D': question.option_d,
                    'correct': question.correct_option
                }

            response_data['questions'].append(question_data)

        return Response(response_data)

    except Exception as e:
        logging.error(f"Error in get_test_for_attempt: {str(e)}", exc_info=True)
        return Response(
            {"error": "Internal server error"},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


class SubmitTestView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, test_id, format=None):
        try:
            print("Received data:", request.data)
            attempt = self._get_or_create_attempt(request.user, test_id)

            if attempt.is_submitted:
                return Response(
                    {"error": "You have already submitted this test."},
                    status=status.HTTP_400_BAD_REQUEST
                )

            answers = self._normalize_answers(request.data.get("answers", {}))
            if not answers:
                return Response({"error": "No answers provided."}, status=400)

            evaluation = self._evaluate_answers(answers, attempt)
            self._finalize_attempt(attempt, evaluation)

            return Response(self._format_results(attempt, evaluation), status=201)

        except Exception as e:
            print("Error in SubmitTestView:", str(e))
            return Response({"error": str(e)}, status=500)

    def _normalize_answers(self, answers):
        if isinstance(answers, list):
            return {str(ans.get('question_id')): ans.get('answer') for ans in answers if 'question_id' in ans and 'answer' in ans}
        elif isinstance(answers, dict):
            return {str(k): v for k, v in answers.items()}
        return {}

    def _get_or_create_attempt(self, user, test_id):
        test = get_object_or_404(Test, pk=test_id)
        total_questions = test.questions.count()
        attempt, _ = TestAttempt.objects.get_or_create(
            student=user,
            test_id=test_id,
            defaults={
                'start_time': timezone.now(),
                'is_submitted': False,
                'total_questions': total_questions
            }
        )
        return attempt

    def _evaluate_answers(self, answers, attempt):
        question_ids = [int(qid) for qid in answers.keys()]
        questions = Question.objects.filter(id__in=question_ids, test_id=attempt.test_id).in_bulk(field_name='id')

        correct_count = 0
        weak_topics = set()
        question_results = []
        ai_input_data = []

        for qid in question_ids:
            if qid not in questions:
                raise ValueError(f"Invalid question ID: {qid}")

            question = questions[qid]
            student_answer = answers[str(qid)].strip()

            is_correct = student_answer.upper() == question.correct_option.upper() if question.question_type == 'MCQ' else False

            if not is_correct:
                ai_input_data.append({
                    'question_id': qid,
                    'question_content': question.content,
                    'question_type': question.question_type,
                    'student_answer': student_answer,
                    'correct_option': question.correct_option,
                })

            if is_correct:
                correct_count += 1

            # Placeholder values, updated later if AI runs
            question_results.append({
                'question_id': qid,
                'student_answer': student_answer,
                'is_correct': is_correct,
                'feedback': "Correct" if is_correct else "Pending...",
                'model_answer': None
            })

            StudentAnswer.objects.update_or_create(
                attempt=attempt,
                question=question,
                defaults={
                    'answer_text': student_answer,
                    'is_correct': is_correct,
                    'ai_feedback': "Correct" if is_correct else "",
                    'suggested_topics': None
                }
            )

        # Run AI only on incorrects
        if ai_input_data:
            ai_feedback_map = self._evaluate_with_AI(ai_input_data)
            for result in question_results:
                qid = result['question_id']
                if qid in ai_feedback_map:
                    feedback_data = ai_feedback_map[qid]
                    result.update(feedback=feedback_data['feedback'], model_answer=feedback_data['model_answer'])
                    weak_topics.update(feedback_data.get('weak_topics', []))

                    # Update in DB
                    StudentAnswer.objects.filter(attempt=attempt, question_id=qid).update(
                        ai_feedback=feedback_data['feedback'],
                        suggested_topics=", ".join(feedback_data.get('weak_topics', []))
                    )

        score = round((correct_count / len(questions)) * 100, 2) if questions else 0.0

        return {
            'score': score,
            'correct_count': correct_count,
            'total_questions': len(questions),
            'weak_topics': list(weak_topics),
            'question_results': question_results
        }


    def _evaluate_with_AI(self, data):
        prompt = """You are an AI tutor. For each question, provide:
    - feedback on why the student's answer is wrong,
    - a model answer (for QNA),
    - a list of weak topics related to the error.

    Return JSON in this format:
    {
    "evaluations": [
        {
        "question_id": 123,
        "feedback": "string",
        "model_answer": "string or null",
        "weak_topics": ["topic1", "topic2"]
        },
        ...
    ]
    }

    Questions:\n""" + json.dumps(data, indent=2)

        try:
            ai_response = evaluate_with_AI(prompt)
            if not ai_response:
                raise ValueError("Empty AI response")

            ai_response = ai_response.strip()
            if ai_response.startswith('```json'):
                ai_response = ai_response[7:-3].strip()
            elif ai_response.startswith('```'):
                ai_response = ai_response[3:-3].strip()

            parsed = json.loads(ai_response)
            return {int(e['question_id']): e for e in parsed.get("evaluations", [])}

        except Exception as e:
            print("AI fallback triggered:", e)
            return self._simple_fallback_evaluation(data)


    def _simple_fallback_evaluation(self, evaluation_data):
        evaluations = []
        for item in evaluation_data:
            if item['type'] == 'MCQ':
                is_correct = item['student_answer'].strip().upper() == item['correct_option']
                feedback = "Correct" if is_correct else f"Correct answer: {item['correct_option']}"
            else:
                is_correct = False
                feedback = "Subjective question - AI feedback not available"
            evaluations.append({
                'question_id': item['question_id'],
                'is_correct': is_correct,
                'feedback': feedback,
                'weak_topics': [],
                'model_answer': None
            })
        return evaluations

    def _finalize_attempt(self, attempt, evaluation):
        attempt.score = evaluation['score']
        attempt.correct_answers = evaluation['correct_count']
        attempt.total_questions = evaluation['total_questions']
        attempt.end_time = timezone.now()
        attempt.is_submitted = True
        attempt.ai_feedback = self._generate_ai_feedback(evaluation)
        attempt.suggested_topics = ", ".join(evaluation['weak_topics'])
        attempt.save()

    def _generate_ai_feedback(self, evaluation):
        if not evaluation['weak_topics']:
            return f"Score: {evaluation['score']}% - Great job!"

        prompt = f"""Generate a 2-3 sentence feedback for the following test result:
- Score: {evaluation['score']}%
- Correct Answers: {evaluation['correct_count']}/{evaluation['total_questions']}
- Weak Topics: {', '.join(evaluation['weak_topics'])}

Feedback:"""
        try:
            return evaluate_with_AI(prompt).strip()
        except:
            return f"Score: {evaluation['score']}%. Focus on: {', '.join(evaluation['weak_topics'])}"

    def _format_results(self, attempt, evaluation):
        return {
            'attempt_id': attempt.id,
            'test_id': attempt.test_id,
            'student': attempt.student.username,
            'score': evaluation['score'],
            'correct_answers': evaluation['correct_count'],
            'total_questions': evaluation['total_questions'],
            'marks': {
                'obtained': evaluation['correct_count'],
                'total': evaluation['total_questions']
            },
            'weak_topics': evaluation['weak_topics'],
            'feedback': attempt.ai_feedback,
            'questions': evaluation['question_results'],
            'submitted_at': attempt.end_time
        }


@method_decorator(csrf_exempt, name='dispatch')
class PracticeGenerateQuestionsView(APIView):
    permission_classes = [IsAuthenticated]
    parser_classes = [MultiPartParser, FormParser]

    def post(self, request):
        try:
            prompt_text = request.data.get('prompt_text', '').strip()
            url = request.data.get('url', '').strip()
            difficulty = request.data.get('difficulty', 'medium')
            file = request.FILES.get('file', None)
            mcq_count = int(request.data.get('mcq_count', 0))
            qna_count = int(request.data.get('qna_count', 0))
            max_tokens = int(request.data.get('max_tokens', 3000))

            if mcq_count + qna_count == 0:
                return Response({"error": "At least one question type must be requested"}, status=400)

            base_text = ""

            if file:
                file_name = default_storage.save(file.name, file)
                file_path = default_storage.path(file_name)
                base_text += self.extract_text_from_file(file_path, file.name)

            if url:
                base_text += "\n" + self.extract_text_from_url(url)

            if prompt_text:
                base_text += "\n" + prompt_text

            if not base_text.strip():
                return Response({"error": "No valid input (text, file, or URL) provided."}, status=400)

            base_text = self.trim_text_by_tokens(base_text, max_tokens)

            questions = []
            if mcq_count > 0:
                questions += self.generate_questions(base_text, difficulty, mcq_count, "mcq")
            if qna_count > 0:
                questions += self.generate_questions(base_text, difficulty, qna_count, "qna")

            return Response({"questions": questions})

        except Exception as e:
            logger.error(f"[PracticeGeneration Error]: {str(e)}")
            return Response({"error": "Failed to generate questions. Please try again."}, status=500)

    def generate_questions(self, prompt, difficulty, count, qtype):
        instruction = {
            "mcq": (
                f"Generate {count} {difficulty} MCQs. Each must include: "
                "'content', 'option_a', 'option_b', 'option_c', 'option_d', and 'correct_option'. "
                "Respond only in JSON as: {\"questions\": [...]}."
            ),
            "qna": (
                f"Generate {count} {difficulty} open-ended questions. "
                "Respond only in JSON as: {\"questions\": [{\"content\": \"...\"}]}."
            )
        }

        full_prompt = (
            f"You are an expert exam question generator.\n"
            f"{instruction[qtype]}\n\n"
            f"Here is the content:\n\"{prompt}\"\n\n"
            f"Respond ONLY in JSON format."
        )

        raw_response = evaluate_with_AI(full_prompt)

        if not raw_response:
            logger.warning("[AI Response] No response received from model.")
            return self.fallback_questions(qtype, count, difficulty)

        try:
            if raw_response.startswith("```json"):
                raw_response = raw_response.split("```json")[1].split("```")[0].strip()
            elif raw_response.startswith("```"):
                raw_response = raw_response.split("```")[1].split("```")[0].strip()

            questions_data = json.loads(raw_response)

            questions = questions_data.get("questions", [])

            for q in questions:
                q['question_type'] = "MCQ" if qtype == "mcq" else "QNA"
                q['difficulty'] = difficulty

            return questions

        except json.JSONDecodeError as json_err:
            logger.warning(f"[AI_Model JSON Decode Error]: {str(json_err)}")
            logger.debug(f"[Raw AI Response]: {repr(raw_response)}")
            return self.fallback_questions(qtype, count, difficulty)

    def fallback_questions(self, qtype, count, difficulty):
        if qtype == "mcq":
            return [{
                'question_type': 'MCQ',
                'content': f'Demo MCQ {i+1}',
                'option_a': 'Option A',
                'option_b': 'Option B',
                'option_c': 'Option C',
                'option_d': 'Option D',
                'correct_option': 'A',
                'difficulty': difficulty
            } for i in range(count)]
        else:
            return [{
                'question_type': 'QNA',
                'content': f'Demo QNA {i+1}',
                'difficulty': difficulty
            } for i in range(count)]

    def extract_text_from_file(self, file_path, filename):
        try:
            if filename.endswith(".pdf"):
                text = ""
                with fitz.open(file_path) as doc:
                    for page in doc:
                        text += page.get_text()
                return text
            elif filename.endswith(".docx"):
                doc = docx.Document(file_path)
                return "\n".join([para.text for para in doc.paragraphs])
            elif filename.endswith(".txt"):
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            else:
                raise ValueError("Unsupported file format")
        except Exception as e:
            logger.error(f"[File Extraction Error]: {str(e)}")
            return ""

    def extract_text_from_url(self, url):
        try:
            parsed_url = urlparse(url)
            if "youtube.com" in parsed_url.netloc or "youtu.be" in parsed_url.netloc:
                return self.extract_text_from_youtube(url)
            else:
                return self.extract_text_from_webpage(url)
        except Exception as e:
            logger.error(f"[URL Extraction Error]: {str(e)}")
            return ""

    def extract_text_from_webpage(self, url):
        try:
            from bs4 import BeautifulSoup
            html = requests.get(url, timeout=10).text
            soup = BeautifulSoup(html, "html.parser")
            text = soup.get_text(separator=" ", strip=True)
            return text
        except Exception as e:
            logger.error(f"[Webpage Extraction Error]: {str(e)}")
            return ""

    def trim_text_by_tokens(self, text, max_tokens=3000):
        words = word_tokenize(text,language='english', preserve_line=True)
        if len(words) > max_tokens:
            words = words[:max_tokens]
        return " ".join(words)

    def extract_text_from_youtube(self, url):
        # Placeholder if you want to support YouTube transcript parsing
        return "YouTube transcript support is not implemented."



class PracticeCheckView(APIView):
    def post(self, request, *args, **kwargs):
        questions = request.data.get('questions', [])
        
        try:
            evaluation_prompt = self.create_evaluation_prompt(questions)
            ai_response = evaluate_with_AI(evaluation_prompt)
            
            if not ai_response:
                return Response({"error": "AI evaluation service unavailable"}, 
                              status=status.HTTP_503_SERVICE_UNAVAILABLE)

            # Extract JSON from possible code block
            json_match = re.search(r'```json\s*({.*?})\s*```', ai_response, re.DOTALL)
            if json_match:
                ai_response = json_match.group(1)
            
            evaluation = json.loads(ai_response)
            return self.format_ai_response(evaluation, questions)
            
        except json.JSONDecodeError as e:
            print(f"JSON Parse Error: {e}\nResponse Content: {ai_response[:500]}")
            return Response({"error": "Invalid AI response format"}, 
                          status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        except Exception as e:
            print(f"Evaluation Error: {str(e)}")
            return Response({"error": "Evaluation failed"}, 
                          status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def create_evaluation_prompt(self, questions):
        prompt = """Act as a computer science expert. Analyze these questions and answers:

        Format your response as pure JSON without Markdown formatting.
        Follow this structure exactly:
        {
            "score_breakdown": {
                "mcq": {"correct": number, "total": number, "percentage": number},
                "qna": {"correct": number, "total": number, "percentage": number}
            },
            "results": [{
                "question": "text",
                "topic": "text",
                "question_type": "MCQ/QNA",
                "options": {
                    "a": "text",
                    "b": "text",
                    "c": "text",
                    "d": "text"
                },
                "student_answer": "text",
                "correct_answer": "text",
                "correct_option": "a-d",
                "is_correct": boolean,
                "marks": number,
                "feedback": "text",
                "key_concepts": ["list"]
            }],
            "overall_feedback": "text",
            "suggested_topics": ["list"]
        }

        Questions and Answers:
        """
        
        for idx, q in enumerate(questions, 1):
            prompt += f"\nQ{idx}: {q.get('content', '')}"
            prompt += f"\nStudent Answer: {q.get('student_answer', '')}"
            if q.get('question_type', '').upper() == 'MCQ':
                prompt += "\nOptions:"
                prompt += f"\nA) {q.get('option_a', '')}"
                prompt += f"\nB) {q.get('option_b', '')}"
                prompt += f"\nC) {q.get('option_c', '')}"
                prompt += f"\nD) {q.get('option_d', '')}"
                prompt += f"\nCorrect Option: {q.get('correct_option', '')}"
            else:
                prompt += f"\nReference Answer: {q.get('correct_answer_text', '')}"
        
        prompt += "\n\nInclude all MCQ options in the response."
        return prompt

    def format_ai_response(self, evaluation, original_questions):
        # Convert scores to proper types
        mcq_correct = 0
        qna_correct = 0
        total_mcq = 0
        total_qna = 0

        # Validate and correct scores based on original questions
        for ai_result, original_q in zip(evaluation['results'], original_questions):
            # Preserve original options for MCQ
            if original_q.get('question_type', '').upper() == 'MCQ':
                # Get original correct option
                original_correct = original_q.get('correct_option', '').lower().replace('option_', '')
                student_answer = ai_result.get('student_answer', '').lower()
                
                # Validate correctness against original data
                actual_correct = student_answer == original_correct
                if ai_result['is_correct'] != actual_correct:
                    print(f"Correcting AI mismatch for question: {original_q.get('content', '')}")
                    ai_result['is_correct'] = actual_correct
                    ai_result['marks'] = 1.0 if actual_correct else 0.0

                # Update counters
                total_mcq += 1
                if actual_correct:
                    mcq_correct += 1

                ai_result['options'] = {
                    'a': original_q.get('option_a', ''),
                    'b': original_q.get('option_b', ''),
                    'c': original_q.get('option_c', ''),
                    'd': original_q.get('option_d', '')
                }
                ai_result['correct_option'] = original_correct
            else:
                # Handle QNA validation if needed
                total_qna += 1
                if ai_result.get('is_correct', False):
                    qna_correct += 1
                ai_result.pop('options', None)
                ai_result.pop('correct_option', None)

            # Type conversions
            ai_result['is_correct'] = bool(ai_result['is_correct'])
            ai_result['marks'] = float(ai_result['marks'])
            ai_result['key_concepts'] = list(ai_result.get('key_concepts', []))

        # Recalculate scores based on validated data
        evaluation['score_breakdown']['mcq'] = {
            'correct': mcq_correct,
            'total': total_mcq,
            'percentage': round((mcq_correct / total_mcq * 100) if total_mcq > 0 else 0, 2)
        }

        evaluation['score_breakdown']['qna'] = {
            'correct': qna_correct,
            'total': total_qna,
            'percentage': round((qna_correct / total_qna * 100) if total_qna > 0 else 0, 2)
        }

        return Response(evaluation, status=status.HTTP_200_OK)

    
class TestResultView(APIView):
    def get(self, request, attempt_id, format=None):
        try:
            attempt = TestAttempt.objects.get(id=attempt_id)
            
            # Verify the requesting user owns this attempt
            if attempt.student != request.user:
                return Response(
                    {"error": "Unauthorized access"},
                    status=status.HTTP_403_FORBIDDEN
                )
            
            student_answers = StudentAnswer.objects.filter(attempt=attempt).select_related('question')

            question_results = []
            weak_topics_set = set()

            for sa in student_answers:
                weak_topics = sa.suggested_topics.split(",") if sa.suggested_topics else []
                weak_topics_set.update([topic.strip() for topic in weak_topics if topic.strip()])
                question_results.append({
                    'question_id': sa.question.id,
                    'content': sa.question.content,
                    'questionType': sa.question.question_type,
                    'options': {
                        'A': sa.question.option_a,
                        'B': sa.question.option_b,
                        'C': sa.question.option_c,
                        'D': sa.question.option_d
                    } if sa.question.question_type == "MCQ" else None,
                    'student_answer': sa.answer_text,
                    'is_correct': sa.is_correct,
                    'feedback': sa.ai_feedback,
                })


            return Response({
                "attempt_id": attempt.id,
                "test_id": attempt.test_id,
                "student": attempt.student.username,
                "score": attempt.score,
                "correct_answers": attempt.correct_answers,
                "total_questions": attempt.total_questions,
                "marks": {
                    "obtained": attempt.correct_answers,
                    "total": attempt.total_questions
                },
                "weak_topics": list(weak_topics_set),
                "feedback": attempt.ai_feedback,
                "questions": question_results,
                "submitted_at": attempt.end_time
            })

        except TestAttempt.DoesNotExist:
            return Response(
                {"error": "Test results not found"},
                status=status.HTTP_404_NOT_FOUND
            )
            
class TeacherSessionTestsView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, session_id):
        if request.user.role != "teacher":
            return Response({"error": "Unauthorized"}, status=403)

        session = get_object_or_404(Session, id=session_id, teacher=request.user)
        serializer = SessionWithTestsSerializer(session)
        return Response(serializer.data)
    
class TeacherTestListView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        tests = Test.objects.filter(teacher=request.user)
        serializer = TestSerializer(tests, many=True)
        return Response(serializer.data)


class TestDetailView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, test_id):
        test = get_object_or_404(Test, id=test_id)
        serializer = TestSerializer(test)
        return Response(serializer.data)

    def patch(self, request, test_id):
        test = get_object_or_404(Test, id=test_id)

        if request.user != test.teacher:
            return Response({'error': 'Permission denied.'}, status=status.HTTP_403_FORBIDDEN)

        # Update Test fields
        serializer = TestSerializer(test, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
        else:
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        # Handle questions
        updated_questions = request.data.get('questions', [])
        deleted_question_ids = request.data.get('delete_questions', [])
        new_questions = request.data.get('new_questions', [])

        # Process updated questions
        for q_data in updated_questions:
            question_id = q_data.get('id')
            if not question_id:
                continue
            try:
                question = test.questions.get(id=question_id)
            except Question.DoesNotExist:
                continue
            serializer = QuestionSerializer(question, data=q_data, partial=True)
            if serializer.is_valid():
                serializer.save()
            else:
                return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        # Process deletions
        for qid in deleted_question_ids:
            try:
                question = test.questions.get(id=qid)
                question.delete()
            except Question.DoesNotExist:
                continue

        # Process new questions - CORRECTED SECTION
        for q_data in new_questions:
            serializer = QuestionSerializer(data=q_data)
            if serializer.is_valid():
                # Set test and teacher directly when saving
                serializer.save(test=test, teacher=request.user)
            else:
                return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        return Response({'message': 'Test updated successfully.'}, status=status.HTTP_200_OK)

    def delete(self, request, test_id):
        test = get_object_or_404(Test, id=test_id)

        if request.user != test.teacher:
            return Response({'error': 'Permission denied.'}, status=status.HTTP_403_FORBIDDEN)

        test.delete()
        return Response({'message': 'Test deleted successfully.'}, status=status.HTTP_204_NO_CONTENT)


class DeleteTeacherSessionDetailView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def delete(self, request, session_id):
        if getattr(request.user, 'role', None) != "teacher":
            return Response({"error": "Only teachers can delete their sessions."}, status=403)

        session = get_object_or_404(Session, id=session_id)

        if session.teacher != request.user:
            return Response({"error": "You do not have permission to delete this session."}, status=403)

        session.delete()
        return Response({"message": "Session deleted successfully."}, status=204)
    

class EnrolledStudentsView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, session_id):
        session = get_object_or_404(Session, id=session_id)

        if request.user != session.teacher:
            return Response({"error": "You are not allowed to view these students."}, status=403)

        students = session.enrolled_students.all()
        serializer = UserSerializer(students, many=True)
        return Response(serializer.data)


'''class StudentReportPDFView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, student_id):
        student = get_object_or_404(User, id=student_id, role='student')

        # Access control: student can see own report, teacher can see enrolled student's report
        if request.user.role == 'student' and request.user != student:
            return HttpResponse("Unauthorized", status=403)

        if request.user.role == 'teacher':
            # Check if the student is enrolled in any session taught by this teacher
            teacher_sessions = Session.objects.filter(teacher=request.user, enrolled_students=student)
            if not teacher_sessions.exists():
                return HttpResponse("Unauthorized", status=403)

        # Fetch attempts for this student
        attempts = TestAttempt.objects.filter(student=student).select_related('test__session', 'test__teacher')

        context = {
            'student': student,
            'attempts': [],
        }

        for attempt in attempts:
            context['attempts'].append({
                'test_title': attempt.test.title,
                'session_name': attempt.test.session.session_name,
                'teacher_name': attempt.test.teacher.get_full_name() or attempt.test.teacher.username,
                'score': attempt.score or attempt.calculate_score(),
                'correct_answers': attempt.correct_answers,
                'total_questions': attempt.total_questions,
                'start_time': attempt.start_time.strftime('%Y-%m-%d %H:%M'),
                'end_time': attempt.end_time.strftime('%Y-%m-%d %H:%M') if attempt.end_time else 'N/A',
                'time_taken': str(attempt.time_taken()),
                'ai_feedback': attempt.ai_feedback,
                'suggested_topics': attempt.suggested_topics or [],
            })

        template = get_template('student_report_template.html')  # create this template
        html = template.render(context)
        pdf_file = BytesIO()
        pisa_status = pisa.CreatePDF(src=html, dest=pdf_file)

        if pisa_status.err:
            return HttpResponse("PDF generation failed", status=500)

        return HttpResponse(pdf_file.getvalue(), content_type='application/pdf')'''



# 1. Get unattempted tests in a session
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def session_detail(request, session_id):
    try:
        session = Session.objects.get(id=session_id)
    except Session.DoesNotExist:
        return Response({'error': 'Session not found'}, status=404)

    all_tests = session.tests.all()
    attempted_test_ids = TestAttempt.objects.filter(student=request.user, test__in=all_tests).values_list('test_id', flat=True)
    
    unattempted_tests = all_tests.exclude(id__in=attempted_test_ids)

    return Response({
        "session": SessionSerializer(session).data,
        "unattempted_tests": TestSerializer(unattempted_tests, many=True).data,
    })


# 2. Get all attempted tests by student
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def student_attempted_tests(request):
    attempts = TestAttempt.objects.filter(student=request.user)
    serializer = AttemptedTestListSerializer(attempts, many=True)
    return Response(serializer.data)

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def attempted_test_detail(request, attempt_id):
    """
    Retrieve detailed info of a single test attempt for the authenticated student.
    Includes per-question feedback, score, weak topics, and submitted answers.
    """
    try:
        attempt = TestAttempt.objects.get(id=attempt_id, student=request.user)
    except TestAttempt.DoesNotExist:
        return Response({'error': 'Test not attempted or not found.'}, status=status.HTTP_404_NOT_FOUND)

    student_answers = StudentAnswer.objects.filter(attempt=attempt).select_related('question')

    question_results = []
    weak_topics_set = set()

    for sa in student_answers:
        weak_topics = sa.suggested_topics.split(",") if sa.suggested_topics else []
        weak_topics_set.update([topic.strip() for topic in weak_topics if topic.strip()])
        
        question_results.append({
            'question_id': sa.question.id,
            'content': sa.question.content,
            'questionType': sa.question.question_type,
            'options': {
                'A': sa.question.option_a,
                'B': sa.question.option_b,
                'C': sa.question.option_c,
                'D': sa.question.option_d
            } if sa.question.question_type == "MCQ" else None,
            'student_answer': sa.answer_text,
            'is_correct': sa.is_correct,
            'feedback': sa.ai_feedback,
        })

    return Response({
        "attempt_id": attempt.id,
        "test_id": attempt.test_id,
        "student": attempt.student.username,
        "score": attempt.score,
        "correct_answers": attempt.correct_answers,
        "total_questions": attempt.total_questions,
        "marks": {
            "obtained": attempt.correct_answers,
            "total": attempt.total_questions
        },
        "weak_topics": list(weak_topics_set),
        "feedback": attempt.ai_feedback,
        "questions": question_results,
        "submitted_at": attempt.end_time
    }, status=status.HTTP_200_OK)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def download_attempt_pdf(request, attempt_id):
    try:
        attempt = TestAttempt.objects.get(id=attempt_id, student=request.user)
    except TestAttempt.DoesNotExist:
        return Response({'error': 'Test attempt not found'}, status=404)

    submitted_at = attempt.submitted_at.strftime('%Y-%m-%d %H:%M') if attempt.submitted_at else "Not Submitted Yet"

    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    
    # Margins and spacing
    left_margin = 50
    right_margin = 50
    top_margin = 50
    line_height = 14
    section_spacing = 20
    y = height - top_margin

    def check_space(required_space=line_height):
        nonlocal y
        if y < (required_space + top_margin):
            p.showPage()
            y = height - top_margin
            return True
        return False

    # Header with styling
    p.setFont("Helvetica-Bold", 16)
    p.setFillColorRGB(0.2, 0.4, 0.6)  # Dark blue color
    p.drawString(left_margin, y, f"Test Report: {attempt.test.title}")
    y -= line_height * 1.5

    # Student information section
    p.setFont("Helvetica", 12)
    p.setFillColorRGB(0, 0, 0)  # Black color
    
    info_lines = [
        f"Student: {request.user.get_full_name() or request.user.username}",
        f"Score: {attempt.score} / {attempt.total_questions}",
        f"Percentage: {round((attempt.score/attempt.total_questions)*100, 2)}%",
        f"Correct Answers: {attempt.correct_answers}",
        f"Submitted At: {submitted_at}"
    ]
    
    for line in info_lines:
        check_space()
        p.drawString(left_margin, y, line)
        y -= line_height

    # AI Feedback section
    y -= section_spacing / 2
    check_space(line_height * 2)
    
    p.setFont("Helvetica-Bold", 12)
    p.drawString(left_margin, y, "AI Feedback:")
    y -= line_height
    
    p.setFont("Helvetica", 11)
    feedback_text = attempt.ai_feedback or 'No feedback available'
    wrapped_feedback = textwrap.wrap(feedback_text, width=90)
    for line in wrapped_feedback:
        check_space()
        p.drawString(left_margin + 10, y, line)
        y -= line_height

    # Suggested Topics section
    if attempt.suggested_topics:
        y -= section_spacing
        check_space(line_height * 3)
        
        p.setFont("Helvetica-Bold", 12)
        p.drawString(left_margin, y, "Suggested Topics for Improvement:")
        y -= line_height * 1.2
        
        try:
            topics = json.loads(attempt.suggested_topics)
            if not isinstance(topics, list):
                topics = attempt.suggested_topics.split(",")
        except json.JSONDecodeError:
            topics = attempt.suggested_topics.split(",")
        
        p.setFont("Helvetica", 11)
        for topic in topics:
            topic = topic.strip()
            if not topic:
                continue
                
            wrapped = textwrap.wrap(f"• {topic}", width=90)
            for line in wrapped:
                check_space()
                p.drawString(left_margin + 15, y, line)
                y -= line_height
            y -= 2  # Small gap between topics

    # Detailed Question Report section
    y -= section_spacing
    check_space(line_height * 3)
    
    p.setFont("Helvetica-Bold", 14)
    p.setFillColorRGB(0.2, 0.4, 0.6)  # Dark blue color
    p.drawString(left_margin, y, "Detailed Question Report")
    y -= line_height * 1.5

    for idx, ans in enumerate(attempt.answers.all(), start=1):
        q = ans.question
        
        # Question header
        check_space(line_height * 4)
        p.setFont("Helvetica-Bold", 12)
        p.setFillColorRGB(0, 0, 0)  # Black color
        question_header = f"Question {idx}"
        p.drawString(left_margin, y, question_header)
        y -= line_height
        
        # Draw a light gray background for the question
        p.setFillColorRGB(0.95, 0.95, 0.95)  # Light gray
        text_height = len(textwrap.wrap(q.content, width=90)) * line_height
        p.rect(left_margin - 5, y - text_height - 5, width - left_margin - right_margin + 10, text_height + 10, fill=1, stroke=0)
        p.setFillColorRGB(0, 0, 0)  # Black color
        
        # Question text
        p.setFont("Helvetica", 11)
        question_text = q.content
        for line in textwrap.wrap(question_text, width=90):
            check_space()
            p.drawString(left_margin, y, line)
            y -= line_height
        
        y -= 5  # Small gap after question
        
        if q.question_type == 'MCQ':
            # MCQ Options
            p.setFont("Helvetica", 10)
            options = [
                f"A. {q.option_a}",
                f"B. {q.option_b}",
                f"C. {q.option_c}",
                f"D. {q.option_d}"
            ]
            
            for opt in options:
                wrapped_opt = textwrap.wrap(opt, width=85)
                for line in wrapped_opt:
                    check_space()
                    p.drawString(left_margin + 10, y, line)
                    y -= line_height
                y -= 2  # Small gap between options
            
            # Answer details
            check_space(line_height * 4)
            p.setFont("Helvetica-Bold", 10)
            p.drawString(left_margin + 10, y, "Your Answer:")
            p.setFont("Helvetica", 10)
            p.drawString(left_margin + 90, y, ans.answer_text)
            y -= line_height
            
            p.setFont("Helvetica-Bold", 10)
            p.drawString(left_margin + 10, y, "Correct Answer:")
            p.setFont("Helvetica", 10)
            p.drawString(left_margin + 90, y, q.correct_option)
            y -= line_height
            
            p.setFont("Helvetica-Bold", 10)
            p.drawString(left_margin + 10, y, "Status:")
            p.setFont("Helvetica", 10)
            status_text = "Correct" if ans.is_correct else "Incorrect"
            p.drawString(left_margin + 90, y, status_text)
            y -= line_height
            
            p.setFont("Helvetica-Bold", 10)
            p.drawString(left_margin + 10, y, "Recommendation:")
            p.setFont("Helvetica", 10)
            rec_text = "Great job!" if ans.is_correct else "Review this concept"
            p.drawString(left_margin + 90, y, rec_text)
            y -= section_spacing / 2

        elif q.question_type == 'QNA':
            # Q&A Answer details
            check_space(line_height * 3)
            p.setFont("Helvetica-Bold", 10)
            p.drawString(left_margin + 10, y, "Your Answer:")
            y -= line_height
            
            p.setFont("Helvetica", 10)
            student_answer = ans.answer_text or "No answer provided"
            wrapped_answer = textwrap.wrap(student_answer, width=90)
            for line in wrapped_answer:
                check_space()
                p.drawString(left_margin + 15, y, line)
                y -= line_height
            
            if ans.ai_feedback:
                y -= 5
                check_space(line_height * 2)
                p.setFont("Helvetica-Bold", 10)
                p.drawString(left_margin + 10, y, "Feedback:")
                y -= line_height
                
                p.setFont("Helvetica", 10)
                wrapped_feedback = textwrap.wrap(ans.ai_feedback, width=90)
                for line in wrapped_feedback:
                    check_space()
                    p.drawString(left_margin + 15, y, line)
                    y -= line_height
            
            y -= section_spacing / 2
        
        # Add a separator line between questions
        check_space(line_height)
        p.line(left_margin, y, width - right_margin, y)
        y -= section_spacing

    # Footer
    p.setFont("Helvetica", 8)
    p.setFillColorRGB(0.5, 0.5, 0.5)  # Gray color
    p.drawString(left_margin, 30, f"Report generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    p.drawRightString(width - right_margin, 30, "Page %d" % p.getPageNumber())

    p.save()
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="Test_Report_{attempt.test.title}_{request.user.username}.pdf"'
    return response




@api_view(['GET'])
@permission_classes([IsAuthenticated])
def teacher_attempt_detail(request, attempt_id):
    try:
        attempt = TestAttempt.objects.select_related('test', 'student').get(id=attempt_id)
    except TestAttempt.DoesNotExist:
        return Response({'error': 'Attempt not found.'}, status=status.HTTP_404_NOT_FOUND)

    # Ensure the authenticated user is the teacher who created the test
    if attempt.test.teacher != request.user:
        return Response({'error': 'Unauthorized access.'}, status=status.HTTP_403_FORBIDDEN)

    student_answers = StudentAnswer.objects.filter(attempt=attempt).select_related('question')

    question_results = []
    weak_topics_set = set()

    for sa in student_answers:
        weak_topics = sa.suggested_topics.split(",") if sa.suggested_topics else []
        weak_topics_set.update([topic.strip() for topic in weak_topics if topic.strip()])
        
        question_results.append({
            'question_id': sa.question.id,
            'content': sa.question.content,
            'questionType': sa.question.question_type,
            'options': {
                'A': sa.question.option_a,
                'B': sa.question.option_b,
                'C': sa.question.option_c,
                'D': sa.question.option_d
            } if sa.question.question_type == "MCQ" else None,
            'student_answer': sa.answer_text,
            'is_correct': sa.is_correct,
            'feedback': sa.ai_feedback,
        })

    return Response({
        "attempt_id": attempt.id,
        "test_id": attempt.test_id,
        "test_title": attempt.test.title,
        "student": {
            "id": attempt.student.id,
            "username": attempt.student.username,
            "email": attempt.student.email
        },
        "score": attempt.score,
        "correct_answers": attempt.correct_answers,
        "total_questions": attempt.total_questions,
        "marks": {
            "obtained": attempt.correct_answers,
            "total": attempt.total_questions
        },
        "weak_topics": list(weak_topics_set),
        "feedback": attempt.ai_feedback,
        "questions": question_results,
        "submitted_at": attempt.end_time
    }, status=status.HTTP_200_OK)



@api_view(['GET'])
@permission_classes([IsAuthenticated])
def teacher_download_attempt_pdf(request, attempt_id):
    try:
        attempt = TestAttempt.objects.select_related('test', 'student').get(id=attempt_id)
    except TestAttempt.DoesNotExist:
        return Response({'error': 'Test attempt not found'}, status=404)

    if attempt.test.teacher != request.user:
        return Response({'error': 'Unauthorized access'}, status=403)

    submitted_at = attempt.submitted_at.strftime('%Y-%m-%d %H:%M') if attempt.submitted_at else "Not Submitted Yet"

    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    y = height - 50

    def check_space(lines=1):
        nonlocal y
        if y < (lines * 20):
            p.showPage()
            y = height - 50

    # Header
    p.setFont("Helvetica-Bold", 14)
    p.drawString(50, y, f"Student Report: {attempt.test.title}")
    y -= 25

    p.setFont("Helvetica", 12)
    p.drawString(50, y, f"Student: {attempt.student.get_full_name() or attempt.student.username}")
    y -= 20
    p.drawString(50, y, f"Email: {attempt.student.email}")
    y -= 20
    p.drawString(50, y, f"Score: {attempt.score} / {attempt.total_questions}")
    y -= 20
    p.drawString(50, y, f"Correct Answers: {attempt.correct_answers}")
    y -= 20
    p.drawString(50, y, f"Submitted At: {submitted_at}")
    y -= 20

    wrapped_feedback = textwrap.wrap(f"AI Feedback: {attempt.ai_feedback or 'Not available'}", width=100)
    for line in wrapped_feedback:
        check_space()
        p.drawString(50, y, line)
        y -= 15

    # Suggested Topics
    if attempt.suggested_topics:
        y -= 10
        check_space()
        p.setFont("Helvetica-Bold", 12)
        p.drawString(50, y, "Suggested Topics for Improvement:")
        y -= 18
        p.setFont("Helvetica", 11)

        try:
            topics = json.loads(attempt.suggested_topics)
            if not isinstance(topics, list):
                topics = attempt.suggested_topics.split(",")
        except json.JSONDecodeError:
            topics = attempt.suggested_topics.split(",")

        for topic in topics:
            wrapped = textwrap.wrap(f"- {topic.strip()}", width=90)
            for line in wrapped:
                check_space()
                p.drawString(60, y, line)
                y -= 13

    y -= 20
    check_space()
    p.setFont("Helvetica-Bold", 12)
    p.drawString(50, y, "Detailed Question Report:")
    y -= 20

    for idx, ans in enumerate(attempt.answers.all(), start=1):
        q = ans.question
        p.setFont("Helvetica-Bold", 11)
        question_text = f"Q{idx}: {q.content}"
        for line in textwrap.wrap(question_text, width=100):
            check_space()
            p.drawString(50, y, line)
            y -= 13

        if q.question_type == 'MCQ':
            p.setFont("Helvetica", 10)
            for opt in ['A', 'B', 'C', 'D']:
                option_text = getattr(q, f'option_{opt.lower()}')
                wrapped_opt = textwrap.wrap(f"{opt}. {option_text}", width=90)
                for line in wrapped_opt:
                    check_space()
                    p.drawString(60, y, line)
                    y -= 12

            check_space(lines=3)
            p.drawString(60, y, f"Student Answer: {ans.answer_text}")
            y -= 12
            p.drawString(60, y, f"Correct Answer: {q.correct_option}")
            y -= 12
            p.drawString(60, y, f"Status: {'✅ Correct' if ans.is_correct else '❌ Incorrect'}")
            y -= 12
            p.drawString(60, y, f"Recommendation: {'Great job!' if ans.is_correct else 'Review related concept'}")
            y -= 18

        elif q.question_type == 'QNA':
            p.setFont("Helvetica", 10)

            student_answer_lines = textwrap.wrap(f"Student Answer: {ans.answer_text}", width=100)
            for line in student_answer_lines:
                check_space()
                p.drawString(60, y, line)
                y -= 12

            if ans.ai_feedback:
                expected_answer_lines = textwrap.wrap(f"Expected Answer: {ans.ai_feedback}", width=100)
                for line in expected_answer_lines:
                    check_space()
                    p.drawString(60, y, line)
                    y -= 12

            check_space()
            p.drawString(60, y, f"Recommendation: {'Good explanation!' if ans.is_correct else 'Work on clarity and depth.'}")
            y -= 18

        check_space()

    p.save()
    buffer.seek(0)
    return HttpResponse(buffer, content_type='application/pdf')



@api_view(['GET'])
@permission_classes([IsAuthenticated])
def teacher_test_attempts(request, test_id):
    try:
        test = Test.objects.get(id=test_id, teacher=request.user)
    except Test.DoesNotExist:
        return Response({'error': 'Test not found or unauthorized'}, status=404)
    
    attempts = TestAttempt.objects.filter(test=test).select_related('student')
    data = [{
        'id': a.id,
        'student_name': a.student.get_full_name() or a.student.username,
        'student_email': a.student.email,
        'correct_answers': a.correct_answers,
        'total_questions': a.total_questions,
        'end_time': a.end_time,
        'score': a.score
    } for a in attempts]
    
    return Response(data)



@api_view(['GET'])
@permission_classes([IsAuthenticated])
def teacher_session_result(request, session_id):
    try:
        session = Session.objects.get(id=session_id, teacher=request.user)
    except Session.DoesNotExist:
        return Response({"error": "Session not found or unauthorized"}, status=status.HTTP_404_NOT_FOUND)

    tests = Test.objects.filter(session=session)
    test_ids = tests.values_list('id', flat=True)

    attempts = TestAttempt.objects.filter(test_id__in=test_ids, is_submitted=True)

    if not attempts.exists():
        return Response({"message": "No attempts found for this session."}, status=status.HTTP_200_OK)

    # Session level stats
    overall_avg = attempts.aggregate(avg_score=Avg('score'))['avg_score']
    highest_score = attempts.aggregate(max_score=Max('score'))['max_score']
    lowest_score = attempts.aggregate(min_score=Min('score'))['min_score']

    # Per student results (only those with student role)
    student_results = []
    students = User.objects.filter(role="student", id__in=attempts.values_list('student_id', flat=True).distinct())

    for student in students:
        student_attempts = attempts.filter(student=student)
        average_score = student_attempts.aggregate(avg=Avg('score'))['avg']
        attempted_tests = student_attempts.count()

        student_results.append({
            "id": student.id,
            "name": student.username,
            "email": student.email,
            "tests_attempted": attempted_tests,
            "average_score": round(average_score, 2) if average_score is not None else None,
        })

    data = {
        "session_name": session.session_name,
        "total_tests": tests.count(),
        "total_students": students.count(),
        "total_attempts": attempts.count(),
        "overall_average_score": round(overall_avg, 2) if overall_avg else None,
        "highest_score": highest_score,
        "lowest_score": lowest_score,
        "students": student_results,
    }

    return Response(data, status=status.HTTP_200_OK)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def teacher_sessions(request):
    sessions = Session.objects.filter(teacher=request.user)
    serializer = SessionSerializer(sessions, many=True)
    return Response(serializer.data)


class SessionDetailView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, session_id):
        session = get_object_or_404(Session, id=session_id)

        # Optional: limit access to teachers who created it or students enrolled
        if request.user.role == 'teacher' and session.teacher != request.user:
            return Response({"error": "Not allowed to view this session"}, status=status.HTTP_403_FORBIDDEN)

        if request.user.role == 'student' and request.user not in session.enrolled_students.all():
            return Response({"error": "You are not enrolled in this session"}, status=status.HTTP_403_FORBIDDEN)

        serializer = SessionSerializer(session)
        return Response(serializer.data)


class NotificationViewSet(viewsets.ViewSet):
    permission_classes = [permissions.IsAuthenticated]
    
    def list(self, request):
        notifications = Notification.objects.filter(
            recipient=request.user
        ).order_by('-created_at')
        serializer = NotificationSerializer(notifications, many=True)
        return Response(serializer.data)
    
    @action(detail=False, methods=['get'])
    def unread(self, request):
        notifications = Notification.objects.filter(
            recipient=request.user,
            is_read=False
        ).order_by('-created_at')
        serializer = NotificationSerializer(notifications, many=True)
        return Response(serializer.data)
    
    @action(detail=True, methods=['patch'])
    def mark_as_read(self, request, pk=None):
        notification = get_object_or_404(
            Notification, 
            pk=pk, 
            recipient=request.user
        )
        notification.mark_as_read()  # Use model method
        return Response({'status': 'marked as read'})
    
    @action(detail=False, methods=['post'])
    def mark_all_read(self, request):
        Notification.objects.filter(
            recipient=request.user,
            is_read=False
        ).update(is_read=True)
        return Response({'status': 'all marked as read'})
    
    @action(detail=False, methods=['get'])
    def count(self, request):
        # FIXED: Use the correct model name
        count = Notification.objects.filter(
            recipient=request.user,
            is_read=False
        ).count()
        return Response({'unread_count': count})
    
class NotificationCountView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        count = Notification.objects.filter(  # Fixed model name
            recipient=request.user,
            is_read=False
        ).count()
        return Response({'unread_count': count})
    
    
@api_view(['POST'])
@permission_classes([IsAuthenticated])
def test_notification(request):
    send_notification(
        recipient=request.user,
        notification_type='test',
        message='Test notification from server',
        content_object=request.user
    )
    return Response({"status": "Test notification sent"})


class UserStatsView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        user = request.user

        if user.role == 'student':
            test_attempts = TestAttempt.objects.filter(student=user, is_submitted=True)
            sessions = Session.objects.filter(enrolled_students=user)


            stats = {
                "role": "student",
                "total_sessions_enrolled": sessions.count(),
                "total_tests_attempted": test_attempts.count(),
                "average_score": round(test_attempts.aggregate(avg=Avg("score"))["avg"] or 0, 2),
            }

        elif user.role == 'teacher':
            sessions = Session.objects.filter(teacher=user)
            tests = Test.objects.filter(teacher=user)
            test_ids = tests.values_list("id", flat=True)
            attempts = TestAttempt.objects.filter(test_id__in=test_ids, is_submitted=True)

            students_count = sessions.values("students").distinct().count()

            stats = {
                "role": "teacher",
                "total_sessions_created": sessions.count(),
                "total_tests_created": tests.count(),
                "total_students_enrolled": students_count,
                "average_student_score": round(attempts.aggregate(avg=Avg("score"))["avg"] or 0, 2),
            }

        else:
            stats = {"detail": "User role not recognized."}

        return Response(stats)

class AdminSessionListView(APIView):
    permission_classes = [IsAuthenticated, IsCustomAdmin]

    def get(self, request):
        sessions = Session.objects.all().select_related("teacher").prefetch_related("tests")
        serializer = AdminSessionSerializer(sessions, many=True)
        return Response(serializer.data)
    
    
class AdminUserListView(ListAPIView):
    serializer_class = UserListSerializer
    permission_classes = [IsAuthenticated, IsCustomAdmin]

    def get_queryset(self):
        role = self.request.query_params.get('role')
        return User.objects.filter(role=role)
    
class AdminUserDeleteView(APIView):
    permission_classes = [IsAuthenticated, IsCustomAdmin]

    def delete(self, request, user_id):
        user = get_object_or_404(User, id=user_id)
        user.delete()
        return Response({"detail": "User deleted successfully"}, status=status.HTTP_204_NO_CONTENT) 